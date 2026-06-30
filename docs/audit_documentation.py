from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


path = Path(__file__).with_name("StudySync_Project_Documentation.docx")
doc = Document(path)
errors = []

if len(doc.sections) != 1:
    errors.append(f"Expected one section, found {len(doc.sections)}")

section = doc.sections[0]
expected = {
    "page_width": 8.5,
    "page_height": 11.0,
    "top_margin": 1.0,
    "right_margin": 1.0,
    "bottom_margin": 1.0,
    "left_margin": 1.0,
}
for attr, inches in expected.items():
    actual = getattr(section, attr).inches
    if abs(actual - inches) > 0.01:
        errors.append(f"{attr}: expected {inches}, found {actual}")

style_expectations = {
    "Normal": (11, 6),
    "Heading 1": (16, 10),
    "Heading 2": (13, 7),
    "Heading 3": (12, 5),
}
for name, (size, after) in style_expectations.items():
    style = doc.styles[name]
    if not style.font.size or abs(style.font.size.pt - size) > 0.1:
        errors.append(f"{name} font size is not {size} pt")
    if not style.paragraph_format.space_after or abs(style.paragraph_format.space_after.pt - after) > 0.1:
        errors.append(f"{name} after spacing is not {after} pt")

numbered_paragraphs = 0
for paragraph in doc.paragraphs:
    if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        numbered_paragraphs += 1
    stripped = paragraph.text.strip()
    if stripped.startswith("• ") or stripped.startswith("- "):
        errors.append(f"Fake bullet detected: {stripped[:60]}")

if numbered_paragraphs < 30:
    errors.append(f"Expected at least 30 real numbered/list paragraphs, found {numbered_paragraphs}")

for index, table in enumerate(doc.tables, start=1):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
    if tbl_w is None or int(tbl_w.get(qn("w:w"))) != 9360:
        errors.append(f"Table {index} does not use 9360 DXA width")
    if tbl_ind is None or int(tbl_ind.get(qn("w:w"))) != 120:
        errors.append(f"Table {index} does not use 120 DXA indent")
    if sum(grid_widths) != 9360:
        errors.append(f"Table {index} grid sums to {sum(grid_widths)}")
    first_row_pr = table.rows[0]._tr.trPr
    if first_row_pr is None or first_row_pr.find(qn("w:tblHeader")) is None:
        errors.append(f"Table {index} has no repeating header marker")
    for row_idx, row in enumerate(table.rows, start=1):
        widths = []
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            widths.append(int(tc_w.get(qn("w:w"))) if tc_w is not None else 0)
        if sum(widths) != 9360:
            errors.append(f"Table {index} row {row_idx} cell widths sum to {sum(widths)}")

with ZipFile(path) as archive:
    bad = archive.testzip()
    if bad:
        errors.append(f"Corrupt ZIP member: {bad}")

if len(doc.inline_shapes) != 4:
    errors.append(f"Expected 4 inline images, found {len(doc.inline_shapes)}")

if errors:
    print("DOCUMENT AUDIT FAILED")
    for error in errors:
        print("-", error)
    raise SystemExit(1)

print(
    f"DOCUMENT AUDIT PASSED | paragraphs={len(doc.paragraphs)} "
    f"tables={len(doc.tables)} lists={numbered_paragraphs} images={len(doc.inline_shapes)}"
)
