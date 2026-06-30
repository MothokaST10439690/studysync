from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUTPUT = DOCS / "StudySync_Project_Documentation.docx"
ASSETS.mkdir(parents=True, exist_ok=True)

# compact_reference_guide with a named StudySync brand override:
# blue heading tokens are consistently replaced by the product copper palette.
INK = "1E1810"
COPPER = "A8642A"
COPPER_LIGHT = "F5EEE7"
COPPER_PALE = "FDF8F2"
BROWN = "6B4A2F"
MUTED = "6B5D50"
LIGHT_MUTED = "A09080"
CREAM = "FAF8F6"
BORDER = "EDE8E2"
WHITE = "FFFFFF"
GREEN = "1A7A40"
RED = "C0392B"
USABLE_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/seguisb.ttf") if bold else Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


F_TITLE = font(43, True)
F_H1 = font(31, True)
F_H2 = font(24, True)
F_BODY = font(20)
F_SMALL = font(16)
F_SMALL_BOLD = font(16, True)


def rounded(draw, box, radius=24, fill=WHITE, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}" if outline else None, width=width)


def draw_arrow(draw, start, end, color=COPPER, width=5):
    draw.line([start, end], fill=f"#{color}", width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) > abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - sign * 18, y2 - 11), (x2 - sign * 18, y2 + 11)]
    else:
        sign = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 11, y2 - sign * 18), (x2 + 11, y2 - sign * 18)]
    draw.polygon(points, fill=f"#{color}")


def create_architecture_illustration(path: Path):
    im = Image.new("RGB", (1600, 880), f"#{CREAM}")
    d = ImageDraw.Draw(im)
    d.text((80, 55), "StudySync application architecture", font=F_H1, fill=f"#{INK}")
    d.text((80, 103), "A compact PHP application with authenticated collaboration services", font=F_BODY, fill=f"#{MUTED}")

    rounded(d, (95, 245, 385, 570), 28, WHITE, BORDER, 3)
    d.text((145, 285), "Student browser", font=F_H2, fill=f"#{INK}")
    for i, label in enumerate(["Responsive pages", "PWA shell", "Upload + download"]):
        y = 365 + i * 55
        d.ellipse((142, y + 5, 158, y + 21), fill=f"#{COPPER}")
        d.text((175, y), label, font=F_BODY, fill=f"#{MUTED}")

    rounded(d, (540, 175, 1040, 650), 32, "161412", None)
    d.text((605, 220), "PHP application", font=F_H1, fill="#FFFFFF")
    boxes = [
        (600, 320, 790, 425, "Pages", "UI + actions"),
        (805, 320, 980, 425, "Auth", "Sessions + CSRF"),
        (600, 455, 790, 560, "Helpers", "Business rules"),
        (805, 455, 980, 560, "Uploads", "Validation"),
    ]
    for x1, y1, x2, y2, title, subtitle in boxes:
        rounded(d, (x1, y1, x2, y2), 18, "26211D", "5B4737", 2)
        d.text((x1 + 20, y1 + 20), title, font=F_SMALL_BOLD, fill=f"#{COPPER}")
        d.text((x1 + 20, y1 + 54), subtitle, font=F_SMALL, fill="#D7CEC5")

    rounded(d, (1190, 170, 1505, 355), 26, WHITE, BORDER, 3)
    d.text((1250, 210), "MySQL", font=F_H2, fill=f"#{INK}")
    d.text((1235, 275), "Users, groups, tasks,", font=F_SMALL, fill=f"#{MUTED}")
    d.text((1235, 305), "messages and tokens", font=F_SMALL, fill=f"#{MUTED}")
    rounded(d, (1190, 470, 1505, 655), 26, WHITE, BORDER, 3)
    d.text((1233, 510), "File storage", font=F_H2, fill=f"#{INK}")
    d.text((1235, 575), "Randomized uploads", font=F_SMALL, fill=f"#{MUTED}")
    d.text((1235, 605), "Membership-gated access", font=F_SMALL, fill=f"#{MUTED}")

    draw_arrow(d, (385, 405), (540, 405))
    draw_arrow(d, (1040, 315), (1190, 275))
    draw_arrow(d, (1040, 510), (1190, 555))
    d.text((418, 365), "HTTPS", font=F_SMALL_BOLD, fill=f"#{COPPER}")
    d.text((1060, 250), "PDO", font=F_SMALL_BOLD, fill=f"#{COPPER}")
    d.text((1055, 545), "stream", font=F_SMALL_BOLD, fill=f"#{COPPER}")

    d.text((80, 810), "Illustration based on the implemented repository structure", font=F_SMALL, fill=f"#{LIGHT_MUTED}")
    im.save(path)


def create_chat_illustration(path: Path):
    im = Image.new("RGB", (1600, 920), f"#{CREAM}")
    d = ImageDraw.Draw(im)
    rounded(d, (60, 55, 1540, 855), 30, WHITE, BORDER, 3)
    d.rectangle((60, 55, 330, 855), fill="#161412")
    d.text((105, 95), "Study", font=F_H2, fill="#FFFFFF")
    d.text((195, 95), "Sync", font=F_H2, fill=f"#{COPPER}")
    d.text((105, 175), "ADVANCED CALCULUS", font=F_SMALL_BOLD, fill="#8E8176")
    for i, label in enumerate(["Tasks", "Members", "Files", "Chat"]):
        y = 245 + i * 66
        if label == "Chat":
            rounded(d, (90, y - 10, 300, y + 43), 12, "34261B", "67462C", 2)
        d.text((125, y), label, font=F_BODY, fill=f"#{COPPER}" if label == "Chat" else "#B0A090")

    d.text((385, 100), "Group chat", font=F_H1, fill=f"#{INK}")
    d.text((385, 150), "Messages and shared files stay in one place", font=F_BODY, fill=f"#{MUTED}")
    rounded(d, (380, 215, 1135, 690), 20, CREAM, BORDER, 2)
    rounded(d, (430, 260, 1035, 365), 18, WHITE, BORDER, 2)
    d.ellipse((455, 285, 505, 335), fill=f"#{COPPER}")
    d.text((472, 294), "N", font=F_SMALL_BOLD, fill="#FFFFFF")
    d.text((530, 277), "Naledi Mokoena", font=F_SMALL_BOLD, fill=f"#{INK}")
    d.text((530, 313), "I cleaned up the revision notes.", font=F_SMALL, fill=f"#{MUTED}")

    rounded(d, (430, 395, 1035, 565), 18, WHITE, BORDER, 2)
    d.rectangle((455, 420, 620, 540), fill=f"#{COPPER_PALE}")
    d.line((480, 500, 585, 455), fill=f"#{COPPER}", width=5)
    d.line((510, 470, 570, 520), fill=f"#{BROWN}", width=5)
    d.text((655, 430), "Calculus revision notes.png", font=F_SMALL_BOLD, fill=f"#{INK}")
    d.text((655, 472), "PNG · 24 KB", font=F_SMALL, fill=f"#{LIGHT_MUTED}")
    rounded(d, (655, 505, 825, 545), 10, COPPER_LIGHT, None)
    d.text((680, 512), "Download", font=F_SMALL_BOLD, fill=f"#{COPPER}")

    d.ellipse((410, 724, 470, 784), fill=f"#{COPPER_LIGHT}", outline=f"#{COPPER}", width=3)
    d.text((429, 724), "+", font=font(38, True), fill=f"#{COPPER}")
    rounded(d, (490, 718, 1035, 792), 17, WHITE, BORDER, 2)
    d.text((525, 742), "Type a message...", font=F_BODY, fill=f"#{LIGHT_MUTED}")
    rounded(d, (1055, 718, 1135, 792), 17, COPPER, None)
    d.text((1074, 742), "Send", font=F_SMALL_BOLD, fill="#FFFFFF")

    rounded(d, (1175, 215, 1485, 690), 20, WHITE, BORDER, 2)
    d.text((1210, 250), "Chat files", font=F_H2, fill=f"#{INK}")
    d.text((1210, 292), "Recent resources", font=F_SMALL, fill=f"#{LIGHT_MUTED}")
    for i, (name, ext) in enumerate([("Revision notes", "PNG"), ("Tutorial set 4", "PDF"), ("Formula sheet", "DOCX")]):
        y = 360 + i * 92
        d.rectangle((1210, y, 1245, y + 44), fill=f"#{COPPER_LIGHT}")
        d.text((1262, y), name, font=F_SMALL_BOLD, fill=f"#{INK}")
        d.text((1262, y + 28), ext, font=F_SMALL, fill=f"#{LIGHT_MUTED}")

    d.text((80, 875), "Interface illustration - actual content depends on group activity", font=F_SMALL, fill=f"#{LIGHT_MUTED}")
    im.save(path)


def create_invite_illustration(path: Path):
    im = Image.new("RGB", (1600, 630), f"#{CREAM}")
    d = ImageDraw.Draw(im)
    d.text((75, 50), "Invitation and notification workflow", font=F_H1, fill=f"#{INK}")
    steps = [
        (75, 175, 360, 455, "1", "Create invite", "A group owner enters an email address."),
        (475, 175, 760, 455, "2", "Share securely", "Registered users see a notification; others receive the link."),
        (875, 175, 1160, 455, "3", "Accept", "The recipient signs in and accepts before expiry."),
        (1275, 175, 1525, 455, "4", "Collaborate", "Membership is added and the owner is notified."),
    ]
    for x1, y1, x2, y2, n, title, body in steps:
        rounded(d, (x1, y1, x2, y2), 26, WHITE, BORDER, 3)
        d.ellipse((x1 + 25, y1 + 25, x1 + 85, y1 + 85), fill=f"#{COPPER}")
        d.text((x1 + 47, y1 + 34), n, font=F_SMALL_BOLD, fill="#FFFFFF")
        d.text((x1 + 28, y1 + 125), title, font=F_H2, fill=f"#{INK}")
        words = body.split()
        lines, current = [], ""
        for word in words:
            candidate = (current + " " + word).strip()
            if d.textlength(candidate, font=F_SMALL) > x2 - x1 - 55:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        for j, line in enumerate(lines):
            d.text((x1 + 28, y1 + 185 + j * 32), line, font=F_SMALL, fill=f"#{MUTED}")
    for a, b in [((360, 315), (475, 315)), ((760, 315), (875, 315)), ((1160, 315), (1275, 315))]:
        draw_arrow(d, a, b)
    d.text((75, 565), "Invitation links expire after seven days and can only be accepted by the intended account.", font=F_BODY, fill=f"#{MUTED}")
    im.save(path)


ARCH_IMG = ASSETS / "architecture.png"
CHAT_IMG = ASSETS / "chat_attachments_illustration.png"
INVITE_IMG = ASSETS / "invitation_workflow.png"
create_architecture_illustration(ARCH_IMG)
create_chat_illustration(CHAT_IMG)
create_invite_illustration(INVITE_IMG)


def set_run_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def apply_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == USABLE_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag_name in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag_name))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_table_header(row)


def set_paragraph_border(paragraph, color=COPPER, size=10, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=LIGHT_MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def create_numbering(doc: Document):
    root = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in root.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in root.findall(qn("w:num"))]

    def add(fmt, marker, left, hanging):
        abstract_id = max(abstract_ids or [0]) + 1
        abstract_ids.append(abstract_id)
        num_id = max(num_ids or [0]) + 1
        num_ids.append(num_id)

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        root.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        root.append(num)
        return num_id

    return add("bullet", "•", 540, 270), add("decimal", "%1.", 540, 270)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.different_first_page_header_footer = True

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, COPPER, 18, 10),
    ("Heading 2", 13, COPPER, 14, 7),
    ("Heading 3", 12, BROWN, 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = rgb(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ("Title", "Subtitle"):
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

styles["Title"].font.size = Pt(30)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = rgb(INK)
styles["Title"].paragraph_format.space_after = Pt(8)
styles["Subtitle"].font.size = Pt(14)
styles["Subtitle"].font.color.rgb = rgb(MUTED)
styles["Subtitle"].paragraph_format.space_after = Pt(18)

bullet_num_id, decimal_num_id = create_numbering(doc)

header = section.header
hp = header.paragraphs[0]
hp.text = "STUDYSYNC  |  PROJECT DOCUMENTATION"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(3)
for run in hp.runs:
    set_run_font(run, size=8.5, bold=True, color=LIGHT_MUTED)
set_paragraph_border(hp, color=BORDER, size=6, space=4)

footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_body(text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        run = p.add_run(text[len(bold_lead):])
        set_run_font(run, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def add_bullet(text):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_number(text):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_callout(label, text, fill=COPPER_PALE, border=COPPER):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8" if edge == "left" else "4")
        node.set(qn("w:space"), "8")
        node.set(qn("w:color"), border)
        p_bdr.append(node)
    p_pr.append(p_bdr)
    lead = p.add_run(label + "  ")
    set_run_font(lead, bold=True, color=border)
    body = p.add_run(text)
    set_run_font(body, color=INK)


def add_table(headers, rows, widths_dxa, header_fill=COPPER_LIGHT, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths_dxa)
    set_table_borders(table)
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_idx % 2 == 1:
                set_cell_shading(cells[i], CREAM)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    apply_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(path: Path, width=6.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption or path.stem.replace("_", " "))
    picture._inline.docPr.set("title", path.stem.replace("_", " ").title())
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(caption)
        set_run_font(r, size=9, italic=True, color=LIGHT_MUTED)


def page_break():
    doc.add_page_break()


# Cover: editorial_cover pattern, restrained StudySync brand treatment.
doc.add_paragraph().paragraph_format.space_after = Pt(70)
icon_path = ROOT / "public" / "icons" / "icon-192.png"
if icon_path.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_icon = p.add_run().add_picture(str(icon_path), width=Inches(1.05))
    cover_icon._inline.docPr.set("descr", "StudySync application icon")
    cover_icon._inline.docPr.set("title", "StudySync")
    p.paragraph_format.space_after = Pt(20)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(12)
r = kicker.add_run("PRODUCT & TECHNICAL DOCUMENTATION")
set_run_font(r, size=10, bold=True, color=COPPER)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("StudySync")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Smart Student Collaboration Platform")

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(8)
rule.paragraph_format.space_after = Pt(34)
set_paragraph_border(rule, color=COPPER, size=14, space=4)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(4)
r = meta.add_run("Version 1.0  |  30 June 2026")
set_run_font(r, size=11, bold=True, color=INK)
meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta2.add_run("Architecture, user guide, deployment and release acceptance")
set_run_font(r, size=10, italic=True, color=MUTED)

doc.add_paragraph().paragraph_format.space_after = Pt(95)
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Prepared for the StudySync project team")
set_run_font(r, size=9.5, color=LIGHT_MUTED)

page_break()

add_heading("Document control", 1)
add_body("This manual describes the current StudySync codebase and the collaboration release added on 30 June 2026. Interface figures are implementation-based illustrations rather than live screenshots.")
add_table(
    ["Field", "Value"],
    [
        ["Document", "StudySync Project Documentation"],
        ["Release", "Collaboration update - version 1.0"],
        ["Audience", "Project owners, developers, testers, administrators and assessors"],
        ["Application stack", "PHP 8.2+, MySQL/MariaDB, responsive HTML/CSS/JavaScript"],
        ["Primary deployment", "Railway/Nixpacks with public as the web document root"],
        ["Database migration", "public/database/2026_06_29_collaboration_features.sql"],
    ],
    [2700, 6660],
)

add_heading("How to use this guide", 2)
add_bullet("Product owners can begin with Sections 1-4 for capabilities and user workflows.")
add_bullet("Developers should use Sections 5-8 for architecture, data, deployment and security controls.")
add_bullet("Testers should use Section 9 as the release acceptance checklist.")
add_bullet("Operations teams should review the production notes before enabling uploads or password email delivery.")

add_heading("Contents", 2)
for item in [
    "1. Executive summary",
    "2. Product overview and user roles",
    "3. Functional feature guide",
    "4. Collaboration workflows",
    "5. Technical architecture",
    "6. Data model and application routes",
    "7. Installation, migration and deployment",
    "8. Security and operations",
    "9. Release acceptance checklist",
    "10. Limitations and recommended roadmap",
]:
    add_bullet(item)

page_break()

add_heading("1. Executive summary", 1)
add_body("StudySync is a compact collaboration workspace for students. It brings groups, tasks, deadlines, shared resources and conversation into one responsive application, reducing the need to coordinate study work across unrelated tools.")
add_callout("Release outcome", "The collaboration update adds chat attachments, an in-chat file panel, secure group invitations, profile pages with photos, password recovery and in-app notifications. The changes are backed by a database migration and updated deployment configuration.")

add_heading("1.1 Release objectives", 2)
add_bullet("Make files part of the conversation, not a separate afterthought.")
add_bullet("Let group owners bring people into private workspaces through controlled invitations.")
add_bullet("Give every student an identifiable profile with useful academic context.")
add_bullet("Provide recoverable accounts without exposing whether an email address exists.")
add_bullet("Surface important group events through a persistent notification centre.")

add_heading("1.2 Capability map", 2)
add_table(
    ["Area", "What StudySync provides", "Primary user value"],
    [
        ["Groups", "Public/private workspaces, join requests and member controls", "A defined home for each study team"],
        ["Tasks", "Due dates, completion status and group filtering", "Clear ownership and deadline visibility"],
        ["Calendar", "Month view of group tasks", "One view of upcoming workload"],
        ["Files", "Validated uploads and member-only downloads", "Shared notes remain available to the group"],
        ["Chat", "Messages, image previews, attachments and recent files", "Conversation and resources stay in context"],
        ["People", "Profiles, invitations and notifications", "Teams are easier to recognise and grow"],
        ["Accounts", "Registration, sign-in and one-time reset links", "Safer, recoverable access"],
    ],
    [1800, 4100, 3460],
    font_size=9,
)

add_heading("2. Product overview and user roles", 1)
add_heading("2.1 Intended users", 2)
add_body("StudySync is designed for students who need a lightweight workspace for coursework, revision groups and shared academic projects. It also supports platform administrators and group creators who manage access and collaboration activity.")
add_table(
    ["Role", "Permissions and responsibilities"],
    [
        ["Student", "Create or join groups, participate in chat, share files, manage assigned tasks, maintain a profile and receive notifications."],
        ["Group creator", "All student capabilities plus join-request review, member removal, invitations and group deletion."],
        ["Administrator", "Platform-level access, group management authority and administrative reporting."],
        ["Invite recipient", "Open a seven-day link, authenticate with the invited account and accept membership."],
    ],
    [2000, 7360],
)

add_heading("2.2 Information architecture", 2)
add_bullet("Dashboard - personal overview of tasks, groups and shared activity.")
add_bullet("Study Groups - discovery, creation and group workspaces.")
add_bullet("Tasks - status, deadline and group filters.")
add_bullet("Files - shared group libraries with upload and download actions.")
add_bullet("Calendar - month-based deadline view.")
add_bullet("Notifications - unread events and navigation to the relevant context.")
add_bullet("Profile - identity, academic information and contribution statistics.")

page_break()

add_heading("3. Functional feature guide", 1)
add_heading("3.1 Group chat and the + attachment control", 2)
add_body("The group Chat tab now supports text-only messages, file-only messages, or a combination of both. The round + control opens the system file picker. After selection, the chosen filename appears beneath the message field before sending.")
add_figure(CHAT_IMG, 6.5, "Figure 1. Implementation-based illustration of chat attachments and the Chat files panel.")
add_bullet("Images render as a bounded preview inside the message and remain downloadable.")
add_bullet("Documents and archives render as compact attachment cards with file type and size.")
add_bullet("The side panel lists the eight most recent group files and links to the complete library.")
add_bullet("A message is accepted when it contains text, an attachment, or both.")

add_heading("3.2 File library and downloads", 2)
add_body("The file library remains available as a dedicated page and as a group tab. Uploads are validated, recorded against a group and uploader, and downloaded through an authenticated endpoint. Direct database identifiers are checked against current membership before any file is streamed.")
add_table(
    ["Category", "Supported formats", "Limit"],
    [
        ["Documents", "PDF, DOCX, PPTX, XLSX, TXT and CSV", "25 MB"],
        ["Archives", "ZIP", "25 MB"],
        ["Images", "PNG, JPG/JPEG, GIF and WebP", "25 MB"],
        ["Profile photo", "PNG, JPG/JPEG and WebP", "5 MB"],
    ],
    [2100, 5160, 2100],
)

add_heading("3.3 Profiles", 2)
add_body("Every active user has a profile page. A user can upload a photo, edit their name, institution, course and biography, and review contribution statistics. Other signed-in members can open a profile from the group member list or from a chat message.")
add_bullet("Activity statistics include joined groups, uploaded files and completed assigned tasks.")
add_bullet("Email addresses are not presented on another student's public profile page.")
add_bullet("The global navigation uses the profile photo when one is available.")

add_heading("3.4 Notifications", 2)
add_body("A global bell displays the unread count. The notification centre lists recent activity, highlights unread items and allows users to mark one or all events as read. Opening an item takes the user to its relevant group or feature page.")
add_bullet("Group invitation created for a registered user.")
add_bullet("Join request submitted to a group creator.")
add_bullet("Join request approved or declined.")
add_bullet("Invitation accepted by a new member.")

page_break()

add_heading("4. Collaboration workflows", 1)
add_heading("4.1 Invite a person to a group", 2)
add_figure(INVITE_IMG, 6.5, "Figure 2. Invitation and in-app notification workflow.")
add_number("Open the group and select the Members tab.")
add_number("Choose Invite people and enter the recipient's email address.")
add_number("Create the invitation. A registered user receives an in-app notification; otherwise copy the generated link.")
add_number("The recipient signs in or registers using the invited email account.")
add_number("The recipient accepts before the seven-day expiry. Membership is inserted once and the group owner is notified.")
add_callout("Access rule", "Only the invited user ID or an account with the invited email address can accept the link. Expired, previously used and mismatched invitations are rejected.")

add_heading("4.2 Recover a forgotten password", 2)
add_number("Choose Forgot password? on the sign-in page and submit the account email address.")
add_number("StudySync creates a random token, stores only its SHA-256 hash and sets a 60-minute expiry.")
add_number("The user opens the reset link, enters a new password of at least eight characters and confirms it.")
add_number("StudySync updates the password hash and marks the token as used so it cannot be replayed.")
add_callout("Production requirement", "Configure APP_URL, MAIL_FROM and a working PHP mail transport or SMTP bridge. Development mode can display the generated link for testing; production mode never does.", fill="FDF6EC", border="B7610A")

add_heading("4.3 Share a file in chat", 2)
add_number("Open a group Chat tab and choose the round + button.")
add_number("Select a supported file no larger than 25 MB.")
add_number("Optionally type a message and choose Send.")
add_number("The upload is saved to the group library and its database ID is linked to the new message.")
add_number("Group members can preview images and download the attachment from the message or Chat files panel.")

add_heading("5. Technical architecture", 1)
add_figure(ARCH_IMG, 6.5, "Figure 3. StudySync logical architecture.")
add_body("The application uses page-level PHP controllers and shared helper modules. Sessions provide authentication state, PDO prepared statements provide database access, and the layout component supplies common navigation and responsive styling.")
add_table(
    ["Layer", "Key files", "Responsibility"],
    [
        ["Presentation", "public/*.php and layout.php", "Page rendering, forms, tabs, modals and responsive navigation"],
        ["Authentication", "config/auth.php and csrf.php", "Session lifecycle, route guards, roles and CSRF verification"],
        ["Business/data", "config/db_helpers.php", "Group, task, file, chat, profile, invitation, notification and reset operations"],
        ["Upload service", "config/uploads.php", "Size/type checks, random filenames, image validation and cleanup"],
        ["Database", "database/init.sql and migration SQL", "Relational data, foreign keys, uniqueness and lifecycle constraints"],
        ["PWA", "manifest.json and service-worker.php", "Install metadata and offline shell behaviour"],
    ],
    [1800, 3000, 4560],
    font_size=8.8,
)

add_heading("5.1 Request path", 2)
add_number("A browser requests a PHP route beneath the public document root.")
add_number("The route loads shared database, authentication and helper modules.")
add_number("Mutating routes validate the session, permissions and a CSRF token.")
add_number("Prepared statements read or change MySQL data; uploads also pass through the upload service.")
add_number("The page captures content and renders it through the shared layout.")

page_break()

add_heading("6. Data model and application routes", 1)
add_heading("6.1 Core tables", 2)
add_table(
    ["Table", "Purpose", "Important relationships"],
    [
        ["users", "Accounts and profile fields", "Referenced by memberships, files, tasks, messages, invitations and notifications"],
        ["study_groups", "Group identity, privacy and creator", "Owns members, tasks, files, messages, requests and invitations"],
        ["group_members", "Many-to-many group membership", "Unique group_id + user_id"],
        ["group_join_requests", "Approval queue for public groups", "Unique pending group/user pair"],
        ["tasks", "Assigned work and due dates", "Belongs to a group and assigned user"],
        ["files", "Upload metadata and storage path", "Belongs to a group and uploader"],
        ["messages", "Group conversation", "Optionally references one file; file deletion sets the reference to null"],
        ["group_invitations", "Expiring, recipient-bound invite links", "Belongs to group, inviter and optional registered user"],
        ["notifications", "Per-user event inbox", "Unread state and optional internal link"],
        ["password_reset_tokens", "One-time password recovery", "Hashed token, expiry and used timestamp"],
    ],
    [2050, 3550, 3760],
    font_size=8.4,
)

add_heading("6.2 New and updated routes", 2)
add_table(
    ["Route", "Audience", "Purpose"],
    [
        ["group.php?tab=chat", "Group members", "Conversation, + attachment upload, previews and recent files"],
        ["group.php?tab=members", "Members/managers", "Profiles, join requests and invitation creation"],
        ["download.php?id=", "Group members", "Membership-gated file stream or image preview"],
        ["profile.php", "Signed-in users", "View profiles and edit the current user's profile"],
        ["invite.php?token=", "Invite recipient", "Validate and accept a group invitation"],
        ["notifications.php", "Signed-in users", "Read event history and clear unread state"],
        ["forgot-password.php", "Signed-out users", "Request a reset link without account enumeration"],
        ["reset-password.php", "Token holder", "Consume a valid token and set a new password"],
    ],
    [2900, 1900, 4560],
    font_size=8.8,
)

add_heading("7. Installation, migration and deployment", 1)
add_heading("7.1 Runtime requirements", 2)
add_bullet("PHP 8.2 or newer with PDO MySQL and mbstring.")
add_bullet("MySQL 8 or MariaDB 10.4 or newer.")
add_bullet("A web server configured with public as its document root.")
add_bullet("Write access to public/uploads for the PHP process.")

add_heading("7.2 Environment variables", 2)
add_table(
    ["Variable", "Required", "Use"],
    [
        ["MYSQLHOST", "Yes", "Database host"],
        ["MYSQLPORT", "Yes", "Database port"],
        ["MYSQLDATABASE", "Yes", "Database name"],
        ["MYSQLUSER", "Yes", "Database user"],
        ["MYSQLPASSWORD", "Yes", "Database password"],
        ["APP_URL", "Recommended", "Canonical base URL for reset links"],
        ["APP_ENV", "Recommended", "Set to production to suppress development reset links"],
        ["MAIL_FROM", "Email delivery", "Sender address used by PHP mail()"],
    ],
    [2450, 1450, 5460],
    font_size=9,
)

add_heading("7.3 Fresh database", 2)
add_number("Create an empty database and a least-privilege database user.")
add_number("Import public/database/init.sql.")
add_number("Set the environment variables and verify the public document root.")
add_number("Ensure the upload directory is writable, then start the application through the chosen server process.")

add_heading("7.4 Existing database upgrade", 2)
add_number("Take and verify a database backup.")
add_number("Run public/database/2026_06_29_collaboration_features.sql exactly once.")
add_number("Deploy the PHP changes after the migration completes.")
add_number("Confirm that messages.file_id and the three new tables are present before opening the site to users.")
add_callout("Ordering matters", "The updated code queries profile and notification columns during sign-in and layout rendering. Apply the migration before deploying the release to an existing database.", fill="FDF1F0", border=RED)

add_heading("7.5 Railway/Nixpacks", 2)
add_body("The repository's Nixpacks, Procfile and nginx configuration now serve /app/public. Set Railway's MySQL variables, APP_URL, APP_ENV=production and mail configuration. The upload directory must be backed by a persistent volume or replaced with object storage because an ephemeral application filesystem can discard user files during redeployments.")

page_break()

add_heading("8. Security and operations", 1)
add_heading("8.1 Implemented controls", 2)
add_table(
    ["Risk", "Implemented control"],
    [
        ["Cross-site request forgery", "CSRF tokens on new account, recovery, profile, invite, notification, file and chat attachment actions"],
        ["Unauthorised file access", "Download route verifies current membership before resolving and streaming a stored path"],
        ["Upload abuse", "Allowlisted extensions, byte-size limits, random stored names and image-content validation"],
        ["Path traversal", "Stored paths must resolve inside public/uploads before reading or deletion"],
        ["Invite misuse", "Random 64-character token, seven-day expiry, pending state and recipient identity/email match"],
        ["Reset token disclosure", "Only a SHA-256 token hash is stored; tokens expire in 60 minutes and are single-use"],
        ["Account enumeration", "Forgot-password response is identical whether or not the address exists"],
        ["Privilege escalation", "New registrations always receive the student role; no hard-coded administrator email remains"],
        ["SQL injection", "User-provided values are passed through PDO prepared statements"],
    ],
    [2600, 6760],
    font_size=8.6,
)

add_heading("8.2 Operational responsibilities", 2)
add_bullet("Store database credentials as platform secrets, never in the repository.")
add_bullet("Use HTTPS in production so sessions, reset links and invitations are protected in transit.")
add_bullet("Back up both MySQL data and persistent file storage.")
add_bullet("Configure a dependable transactional-email provider for password recovery.")
add_bullet("Review upload storage growth and establish retention or quota policies.")
add_bullet("Keep PHP, MySQL and third-party CDN dependencies patched.")

add_heading("8.3 Privacy guidance", 2)
add_body("Profiles are visible to authenticated users who encounter them through group and chat contexts. Avoid placing private contact details in the biography. Administrators should define retention expectations for messages, uploads, notifications and expired tokens.")

add_heading("9. Release acceptance checklist", 1)
add_body("Run this checklist against a staging database after applying the migration. Record evidence and defects in the final two columns.")
add_table(
    ["Area", "Acceptance check", "Result", "Evidence / notes"],
    [
        ["Authentication", "Register, sign in, sign out and recover a password with a one-time token", "Pending", ""],
        ["Profile", "Update name, academic fields, bio and a valid profile image", "Pending", ""],
        ["Invite", "Create, copy, notify, accept and reject an expired/mismatched invitation", "Pending", ""],
        ["Chat", "Send text-only, file-only, image and combined messages", "Pending", ""],
        ["Files", "Download as a member; confirm a non-member receives no file", "Pending", ""],
        ["Notifications", "Unread badge, single open and mark-all-read work correctly", "Pending", ""],
        ["Groups", "Join request, approval, decline and member removal generate expected state", "Pending", ""],
        ["Responsive UI", "Verify desktop, tablet and phone layouts including the chat file panel", "Pending", ""],
        ["Persistence", "Redeploy without losing database records or uploaded files", "Pending", ""],
        ["Regression", "Dashboard, tasks, calendar, admin and PWA install remain functional", "Pending", ""],
    ],
    [1500, 4700, 1100, 2060],
    font_size=8.1,
)

add_heading("9.1 Static verification completed", 2)
add_bullet("Every PHP file passed PHP 8.2 syntax lint after the release changes.")
add_bullet("The repository diff passed git's whitespace error check.")
add_bullet("The database model includes foreign keys and indexes for all new collaboration tables.")
add_callout("Scope note", "Application startup and deployment are intentionally left to the project owner. Complete the staging checklist before treating the release as production-ready.")

add_heading("10. Limitations and recommended roadmap", 1)
add_heading("10.1 Current limitations", 2)
add_bullet("PHP mail() depends on server configuration and is not a complete transactional-email solution by itself.")
add_bullet("Uploaded files are stored on the application filesystem; Railway deployments need persistent storage or object storage.")
add_bullet("Notifications are in-app only; browser push notifications are not implemented.")
add_bullet("Chat uses page refresh rather than WebSockets or live polling.")
add_bullet("Invitation delivery is in-app for existing users and link-based for everyone else.")

add_heading("10.2 Recommended next increments", 2)
add_number("Move uploads to S3-compatible object storage and add per-group quotas.")
add_number("Integrate a transactional-email provider for invitations and password recovery.")
add_number("Add near-real-time chat updates and notification polling.")
add_number("Add antivirus scanning, content-type verification and configurable retention for uploads.")
add_number("Add automated integration tests for role boundaries, token expiry and authenticated downloads.")
add_number("Add an administrator-controlled role-management workflow with an auditable change log.")

add_callout("Definition of done", "The collaboration release is ready for owner-led staging when the database migration is applied, production secrets are configured, persistent upload storage is selected and every acceptance check in Section 9 has evidence.")

# Final preset audit metadata and document properties.
doc.core_properties.title = "StudySync Project Documentation"
doc.core_properties.subject = "Product, architecture, deployment and release guide"
doc.core_properties.author = "StudySync Project Team"
doc.core_properties.keywords = "StudySync, PHP, MySQL, documentation, collaboration"
doc.core_properties.comments = "Generated from the StudySync repository on 30 June 2026."

doc.save(OUTPUT)
print(OUTPUT)
